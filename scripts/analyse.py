"""
Data Analysis and Statistical Utilities

This script provides tools for analyzing datasets, calculating statistical metrics, and generating detailed reports. It includes functionality for descriptive statistics, hypothesis testing, and inequality indices, as well as tools for formatting results into a readable document format (e.g., DOCX).

Key functionalities include:
- Perfroming the mean and proportional differences
- Performing the regression
- Performing the analysis for missing data
- Reporting the result for descriptive and analytical analysis

Dependencies:
- Utilities: pandas, numpy, tabulate, python-docx
- Statistics: scipy, statsmodels, pingouin
- Visualization: matplotlib, seaborn

Usage:
1. Import the functions into your main script or notebook.
2. ...

Authors: P Sitthirat et al
Version: 1.0
License: MIT License
"""

# Utility library imports
from tabulate import tabulate
from docx import Document
import pandas as pd
import numpy as np
from pandas.api.types import is_numeric_dtype

# Statistic library imports
import scipy.stats as stats
from statsmodels.stats.anova import AnovaRM
from statsmodels.stats.outliers_influence import variance_inflation_factor
import statsmodels.api as sm
import statsmodels.formula.api as smf
import pingouin as pg

# Visualization library imports
import matplotlib.pyplot as plt
import seaborn as sns

class ResultExport:
    """
    A class for export the result to docx.
    """
    
    @staticmethod
    def add_to_docx(results, table_name, output_dir):
        """
        Export analysis results to a Word document.

        Parameters:
        - results (DataFrame or list): Analysis results to include in the document.
        - table_name (str): Title of the table in the document.
        - output_dir (str): Directory to save the document.
        
        Returns:
        - None
        """
        
        doc = Document()
        
        # Add a title for the table
        doc.add_heading(table_name, level=2)
        
        # Check if results is a DataFrame or list and if it's empty
        if isinstance(results, pd.DataFrame):
            if results.empty:
                doc.add_paragraph("No data available.")
            else:
                # Convert DataFrame to list of dictionaries
                results = results.to_dict(orient='records')
        elif not results:
            doc.add_paragraph("No data available.")
            return

        # Add a table with a header row if there is data
        if results:
            headers = results[0].keys()
            table = doc.add_table(rows=1, cols=len(headers))
            hdr_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                hdr_cells[i].text = header

            # Add the data rows
            for row_data in results:
                row_cells = table.add_row().cells
                for i, header in enumerate(headers):
                    row_cells[i].text = str(row_data[header])
        
        # Save the document
        doc.save(f'{output_dir}/{table_name}.docx')

class Difference:
    """
    A class for conduct the different analysis among group based on their distribution (parametric / non-parametric test)
    """
    
    @staticmethod
    def normality(df, factor, group_var):
        
        normality_pvalues = {}
        for group in df[group_var].dropna().unique():
            group_data = df[df[group_var] == group][factor].dropna()
            if len(group_data) > 5000:
                ad_stat, critical_values, _ = stats.anderson(group_data)
                pvalue = (ad_stat > critical_values[-1])
                normality_pvalues[group] = 0 if pvalue else 1
            else:
                if len(group_data) > 3:
                    stat, pvalue = stats.shapiro(group_data)
                    normality_pvalues[group] = pvalue
        
        is_normal = all(p > 0.05 for p in normality_pvalues.values())
        
        return is_normal
    
    @staticmethod
    def hov(df, factor, group_var):
        """
        Tests the homogeneity of variances (HOV) across groups.
        
        Parameters:
        df : pandas.DataFrame
            The dataframe containing the data.
        factor : str
            The name of the column to test for homogeneity of variances.
        group_var : str
            The name of the column representing groups.

        Returns:
        bool
            True if variances are homogeneous (p-value > 0.05), False otherwise.
        """
        # Extract groups based on group_var
        grouped_data = [df.loc[df[group_var] == group, factor].dropna() for group in df[group_var].unique()]
        
        # Perform Levene's test
        stat, p_value = stats.levene(*grouped_data)
        
        # Check if p-value is above the threshold for significance
        is_hov = p_value > 0.05
        
        return is_hov

    @staticmethod
    def sphericity(df, subject_col, time_col, value_col):
        """
        Checks the sphericity of the dataset using Mauchly's Test.

        Parameters:
        df (pd.DataFrame): DataFrame in long format with subject ID, time point, and measurement.
        subject_col (str): Name of the subject identifier column.
        time_col (str): Name of the time point (within-subject factor) column.
        value_col (str): Name of the dependent variable (measurement) column.

        Returns:
        dict: Dictionary with Mauchly's test result and sphericity status.
        """
        try:
            # Pivot data to wide format for Mauchly's test
            wide_data = df.pivot(index=subject_col, columns=time_col, values=value_col)
            
            # Perform Mauchly's test
            mauchly_result = pg.mauchly(wide_data)
            
            # Extract test results
            w_stat = mauchly_result['W'].iloc[0]
            chi2 = mauchly_result['chi2'].iloc[0]
            pval = mauchly_result['pval'].iloc[0]
            is_spherical = pval >= 0.05  # True if sphericity holds
            
            return is_spherical
        except Exception as e:
            print(f"Error in is_spherical function: {e}")
            return None

    @staticmethod
    def interpret_effect_size(effect_size, measure):
        """Interpret the effect size based on the measure."""
        if measure == 'd':
            if abs(effect_size) < 0.2:
                return "Negligible"
            elif abs(effect_size) < 0.5:
                return "Small"
            elif abs(effect_size) < 0.8:
                return "Medium"
            else:
                return "Large"
        elif measure == 'n':
            if effect_size < 0.01:
                return "Negligible"
            elif effect_size < 0.06:
                return "Small"
            elif effect_size < 0.14:
                return "Medium"
            else:
                return "Large"
        elif measure == 'r':
            if effect_size < 0.1:
                return "Negligible"
            elif effect_size < 0.3:
                return "Small"
            elif effect_size < 0.5:
                return "Medium"
            else:
                return "Large"
        elif measure == 'v':  # Cramér's V
            if effect_size < 0.1:
                return "Negligible"
            elif effect_size < 0.3:
                return "Small"
            elif effect_size < 0.5:
                return "Medium"
            else:
                return "Large"
        elif measure == 'corr':
            if abs(effect_size) < 0.2:
                return "Negligible"
            elif abs(effect_size) < 0.3:
                return "Small"
            elif abs(effect_size) < 0.4:
                return "Medium"
            elif abs(effect_size) < 0.7:
                return "Large"
            else:
                return "Very Large"     
        elif measure == 'odd':
            if abs(effect_size) < 0.1:
                return "Negligible"
            elif abs(effect_size) < 0.4:
                return "Small"
            elif abs(effect_size) < 1.1:
                return "Medium"
            else:
                return "Large"              
        else:
            return(None)
    
    @staticmethod
    def numerical(df, factor, group_var, subject_id=None, independent=False):
        if pd.notna(subject_id):
            df_test = df[[factor, group_var, subject_id]].dropna()
        else:
            df_test = df[[factor, group_var]].dropna()
        is_normal = Difference.normality(df_test, factor, group_var)
        is_hov = Difference.hov(df_test, factor, group_var)
        groups = [df_test[df_test[group_var] == group][factor].dropna() for group in df_test[group_var].dropna().unique()]

        test = None
        pvalue = None  # Initialize p_value
        eff = None  # Initialize effect size
        interpretation = None  # Initialize interpretation
        
        if is_normal:
            if independent:
                if is_hov: # For homogeniety of variances
                    if len(groups) == 2 and all(len(group) > 1 for group in groups):
                        # Perform independent t-test
                        _, pvalue = stats.ttest_ind(*groups, equal_var=True)
                        
                        diff = np.mean(groups[0]) - np.mean(groups[1])
                        n1, n2 = len(groups[0]), len(groups[1])
                        s1, s2 = np.var(groups[0], ddof=1), np.var(groups[1], ddof=1)
                        pooled_std = np.sqrt(((n1 - 1) * s1 + (n2 - 1) * s2) / (n1 + n2 - 2))
                        cohen_d = diff / pooled_std
                        hedges_g = cohen_d * (1 - (3 / (4 * (n1 + n2) - 9)))
                        effect_size = hedges_g if n1 < 20 or n2 < 20 else cohen_d
                        test = "Independent t-Test"
                        interpretation = Difference.interpret_effect_size(effect_size, 'd')
                    elif len(groups) > 2 and all(len(group) > 1 for group in groups):
                        # Perform One-way ANOVA
                        _, pvalue = stats.f_oneway(*groups)
                            
                        n_total = sum(len(group) for group in groups)
                        n1, n2 = len(groups[0]), len(groups[1])
                        grand_mean = np.mean(np.hstack(groups))
                        ss_between = sum([len(group) * (np.mean(group) - grand_mean) ** 2 for group in groups])
                        ss_within = sum([sum((val - np.mean(group)) ** 2 for val in group) for group in groups])
                        ss_total = ss_between + ss_within
                        df_between = len(groups) - 1
                        df_within = n_total - len(groups)
                        ms_error = ss_within / df_within
                        eta_squared = ss_between / ss_total
                        omega_squared = (ss_between - (df_between * ms_error)) / (ss_total + ms_error)
                        effect_size = omega_squared if n1 < 20 or n2 < 20 else eta_squared
                        test = "One-way ANOVA"
                        interpretation = Difference.interpret_effect_size(effect_size, 'n')
                else:
                    if len(groups) == 2 and all(len(group) > 1 for group in groups):
                        # Perform Welch's t-test
                        _, pvalue = stats.ttest_ind(*groups, equal_var=False)
                        
                        diff = np.mean(groups[0]) - np.mean(groups[1])
                        s1, s2 = np.var(groups[0], ddof=1), np.var(groups[1], ddof=1)
                        pooled_std = np.sqrt((s1+s2)/2)
                        cohen_d = diff / pooled_std
                        hedges_g = cohen_d * (1 - (3 / (4 * (n1 + n2) - 9)))
                        effect_size = hedges_g if n1 < 20 or n2 < 20 else cohen_d
                        test = "Welch's t-Test"
                        interpretation = Difference.interpret_effect_size(effect_size, 'd')
                    elif len(groups) > 2 and all(len(group) > 1 for group in groups):
                        # Perform Welch's ANOVA
                        result = pg.welch_anova(dv=factor, between=group_var, data=df)
                        pvalue = result['p-unc'].iloc[0]
                        
                        n_total = sum(len(group) for group in groups)
                        grand_mean = np.mean(np.hstack(groups))
                        ss_between = sum([len(group) * (np.mean(group) - grand_mean) ** 2 for group in groups])
                        ss_within = sum([sum((val - np.mean(group)) ** 2 for val in group) for group in groups])
                        ss_total = ss_between + ss_within
                        df_between = len(groups) - 1
                        df_within = n_total - len(groups)
                        ms_error = ss_within / df_within
                        omega_squared = (ss_between - (df_between * ms_error)) / (ss_total + ms_error)
                        effect_size = omega_squared
                        test = "Welch's ANOVA"
                        interpretation = Difference.interpret_effect_size(effect_size, 'n')
            else:
                if len(groups) == 2 and all(len(group) > 1 for group in groups):
                    # Perform paired t-Test
                    _, pvalue = stats.ttest_rel(*groups)
                    
                    differences = groups[0] - groups[1]
                    mean_diff = np.mean(differences)
                    std_diff = np.std(differences, ddof=1)
                    cohen_dz = mean_diff / std_diff
                    hedges_gz = cohen_dz * (1 - (3 / (4 * len(differences) - 1)))
                    effect_size = hedges_gz if n1 < 20 or n2 < 20 else cohen_dz
                    test = "Paired t-Test"
                    interpretation = Difference.interpret_effect_size(effect_size, 'd')
                elif len(groups) > 2 and all(len(group) > 1 for group in groups):
                    is_spherical = Difference.sphericity(df_test, subject_col=subject_id, time_col=group_var, value_col=factor)
                    if is_spherical:
                        # Perform repeated measure ANOVA
                        rm_anova = AnovaRM(data=df_test, depvar=factor, subject=subject_id, within=[group_var])
                        rm_result = rm_anova.fit()
                        pvalue = rm_result.anova_table['Pr > F'][0]
                        
                        anova_table = rm_result.anova_table
                        ss_effect = anova_table.loc[anova_table.index[0], 'Sum Sq']
                        ss_error = anova_table.loc[anova_table.index[0], 'Error']
                        partial_eta_squared = ss_effect / (ss_effect + ss_error)
                        effect_size = partial_eta_squared
                        test = "Repeated measure ANOVA"
                        interpretation = Difference.interpret_effect_size(effect_size, 'n')
        else: 
            if independent:
                if len(groups) == 2 and all(len(group) > 1 for group in groups):
                    # Perform Mann-Whitney U test
                    _, pvalue = stats.mannwhitneyu(*groups)
                    
                    n1, n2 = len(groups[0]), len(groups[1])
                    u_stat = min(stats.mannwhitneyu(*groups).statistic, stats.mannwhitneyu(groups[1], groups[0]).statistic)
                    r = 1 - (2 * u_stat / (n1 * n2))
                    effect_size = r
                    test = "Mann-Whitney U Test"
                    interpretation = Difference.interpret_effect_size(effect_size, 'r')
                elif len(groups) > 2 and all(len(group) > 1 for group in groups):
                    # Perform Kruskal-Willis Test
                    h_stat, pvalue = stats.kruskal(*groups)
                    
                    num_groups = len(groups)
                    total_samples = sum(len(group) for group in groups)
                    eta_squared = (h_stat - num_groups + 1) / (total_samples - num_groups)
                    effect_size = eta_squared   
                    test = "Kruskal-Willis test"
                    interpretation = Difference.interpret_effect_size(effect_size, 'n')
            else:
                if len(groups) == 2 and all(len(group) > 1 for group in groups):
                    # Perform Wilcoxon Signed-Rank Test
                    stat, pvalue = stats.wilcoxon(*groups)
                    
                    differences = np.array(groups[0]) - np.array(groups[1])
                    non_zero_differences = differences[differences != 0]
                    n = len(non_zero_differences)
                    z = stat - (n * (n + 1) / 4)  # Adjust Wilcoxon stat to mean
                    sigma = np.sqrt(n * (n + 1) * (2 * n + 1) / 24)  # Standard deviation
                    z = z / sigma  # Standardized Z
                    r = z / np.sqrt(n)
                    effect_size = r
                    test = "Wilcoxon Signed-Rank test"
                    interpretation = Difference.interpret_effect_size(effect_size, 'r')
                elif len(groups) > 2 and all(len(group) > 1 for group in groups):
                    # Perform Friedman Test
                    q_stat, pvalue = stats.friedmanchisquare(*groups)
                    
                    data = np.array(df_test)
                    n, k = data.shape
                    kendalls_w = q_stat / (k * (n * (k + 1) / 2))
                    effect_size = kendalls_w
                    test = "Friedman test"
                    interpretation = Difference.interpret_effect_size(effect_size, 'r')

        return pvalue, effect_size, test, interpretation
      
    @staticmethod
    def proportional(df, factor, group_var):
        
        df_test = df[[factor, group_var]].copy()
        df_test = df_test.dropna()
        
        # Create the contingency table
        contingency_table = pd.crosstab(df_test[factor], df_test[group_var])

        # Perform the Chi-square test
        chi2_stat, pvalue, dof, expected = stats.chi2_contingency(contingency_table)

        # Calculate Cramér's V for effect size
        n = contingency_table.sum().sum()  # Total number of observations
        min_dim = min(contingency_table.shape) - 1  # Minimum of (rows - 1, columns - 1)
        effect_size = np.sqrt(chi2_stat / (n * min_dim))

        # Interpret the effect size
        interpretation = Difference.interpret_effect_size(effect_size, 'v')

        # Test name
        test = "Chi-square Test of Independence"
        
        return pvalue, effect_size, test, interpretation
    
    @staticmethod  
    def correlation(df, var_1, var_2):
        df_test = df[[var_1, var_2]].dropna()
        corr, pvalue = stats.spearmanr(df_test[var_1], df_test[var_2])
        interpretation = Difference.interpret_effect_size(corr, 'corr')
        return corr, pvalue, interpretation
        
class Missing:
    """
    A class for missing analysis
    """

    @staticmethod
    def missing_visualize(df, df_name=None):
        """
        Visualize missing data using heatmap.

        Parameters:
        - df (DataFrame): The input dataset for analyse.
        """
        
        plt.figure(figsize=(12, 6))
        sns.heatmap(df.isnull(), cbar=False, cmap="viridis")
        plt.title(f"Missing Data Heatmap: {df_name}")
        plt.show()

    @staticmethod
    def mcar(df_input, alpha=0.05):
        """
        Test that the missing data is Missing Completely at Random (MCAR)
        
        Parameters:
        - df (DataFrame): The input dataset for analyse.
        """

        df = df_input.copy()
        df.columns = ['x' + str(i) for i in range(df.shape[1])]
        df['missing'] = np.sum(df.isnull(), axis=1)
        n = df.shape[0]
        k = df.shape[1] - 1
        f = k * (k - 1) / 2
        chi2_crit = stats.chi2.ppf(1 - alpha, f)
        chi2_val = ((n - 1 - (k - 1) / 2) ** 2) / (k - 1) / ((n - k) * np.mean(df['missing']))
        p_val = 1 - stats.chi2.cdf(chi2_val, f)
        if chi2_val > chi2_crit:
            print(
                'Reject null hypothesis: Data is not MCAR (p-value={:.4f}, chi-square={:.4f})'.format(p_val, chi2_val)
            )
        else:
            print(
                'Do not reject null hypothesis: Data is MCAR (p-value={:.4f}, chi-square={:.4f})'.format(p_val, chi2_val)
            )

class Descriptive:
    """
    A class for descriptive analysis
    - Reporting the descriptive analysis of categorical and numerical variables.
    """
    
    @staticmethod
    def des_cat(df_input, factor, group_var=None, p_value=None, eff=None, intp=None):
        rows = []
        
        if group_var:
            df = df_input.dropna(subset=[group_var])

        # Header row
        header_row = {'Characteristics': factor, 'Total': '', 'P-value': p_value, 'Effect Size': eff, 'Interpretation': intp}
        if group_var:
            groups = df[group_var].dropna().unique()
            for group in groups:
                header_row[group] = ''
        rows.append(header_row)

        # Calculate total for each category
        categories = sorted(df[factor].dropna().astype(str).unique())
        grand_total = len(df[factor].dropna())

        for category in categories:
            subtotal = len(df[df[factor] == category])
            percent = (subtotal / grand_total) * 100
            row = {'Characteristics': category, 'Total': f'{subtotal:,}\n({percent:.2f}%)'}

            if group_var:
                for group in groups:
                    group_total = len(df[(df[group_var] == group) & (df[factor].notna())])
                    group_count = len(df[(df[factor] == category) & (df[group_var] == group)])
                    group_percent = (group_count / group_total) * 100
                    row[group] = f'{group_count:,}\n({group_percent:.2f}%)'

            row['P-value'] = ''
            row['Effect Size'] = ''
            row['Interpretation'] = ''
            rows.append(row)

        return rows

    @staticmethod
    def des_num(df, factor, group_var=None, p_value=None, eff=None, intp=None):
        rows = {}

        is_normal = Difference.normality(df, factor, group_var)
        
        if is_normal:
            mean = df[factor].dropna().mean()
            sd = df[factor].dropna().std()
            rows = {
                'Characteristics': factor,
                'Total': f'{mean:.2f}\n({sd:.2f})',
                'P-value': p_value,
                'Effect Size': eff, 
                'Interpretation': intp
            }

            if group_var:
                groups = sorted(df[group_var].dropna().unique())
                for group in groups:
                    group_df = df[df[group_var] == group]
                    submean = group_df[factor].dropna().mean()
                    subsd = group_df[factor].dropna().std()
                    rows[group] = f'{submean:.2f}\n({subsd:.2f})'

        else:
            median = df[factor].dropna().median()
            p25 = df[factor].dropna().quantile(0.25)
            p75 = df[factor].dropna().quantile(0.75)
            rows = {
                'Characteristics': factor,
                'Total': f'{median:.2f}\n({p25:.2f}-{p75:.2f})',
                'P-value': p_value,
                'Effect Size': eff, 
                'Interpretation': intp
            }

            if group_var:
                groups = sorted(df[group_var].dropna().unique())
                for group in groups:
                    group_df = df[df[group_var] == group]
                    submedian = group_df[factor].dropna().median()
                    subp25 = group_df[factor].dropna().quantile(0.25)
                    subp75 = group_df[factor].dropna().quantile(0.75)
                    rows[group] = f'{submedian:.2f}\n({subp25:.2f}-{subp75:.2f})'

        return rows

    @staticmethod
    def describe(df, factors, group_var, overall=False, table_name=None, export_result=False):
        
        print(f"{table_name}")
        
        df_test = df.copy()
        results = pd.DataFrame(columns=["Characteristics", "Total"] + list(df[group_var].dropna().unique()) + ["P-value", "Effect Size", "Interpretation"])

        for factor in factors:
            if df_test[factor].dtype == 'O':
                p_val, eff, _, intp = Difference.proportional(df_test, factor, group_var)
                p_value = f'{p_val:.2f}'
                effect_size = f'{eff:.2f}'
                des = Descriptive.des_cat(df_test, factor, group_var, p_value, effect_size, intp)
                descriptive_df = pd.DataFrame(des)

            elif is_numeric_dtype(df[factor]):
                
                p_val, eff, test, intp = Difference.numerical(df_test, factor, group_var, independent=True)
                print(f'Analysis of {factor} using {test}')
                p_value = f'{p_val:.2f}'
                effect_size = f'{eff:.2f}'
                des = Descriptive.des_num(df, factor, group_var, p_value, effect_size, intp)
                descriptive_df = pd.DataFrame([des])

            results = pd.concat([results, descriptive_df], ignore_index=True)

        if not overall:
            results.drop(columns=['Total'], inplace=True)

        print(tabulate(results, showindex=False, headers="keys"))
        if export_result:
            ResultExport.add_to_docx(results, table_name, output_dir='output/analyse')

       
class Regression:
    """
    A class for regression analysis
    """

    @staticmethod
    def binary_logistic(df, independent_var, dependent_var, independent_assign=None, dependent_assign=None, table_name=None, export_result=False, vif_threshold = 5):
        
        # Drop missing values
        df_test = df[[dependent_var] + independent_var].dropna()

        # Ensure the dependent variable is categorical
        df_test[dependent_var] = df_test[dependent_var].astype('category')
        
        # Conduct colinearity testing
        cat_vars = df_test[independent_var].select_dtypes(include=['object', 'category']).columns.tolist()
        num_vars = df_test[independent_var].select_dtypes(include=['number', 'bool']).columns.tolist()
        df_dummy = pd.get_dummies(df_test, columns=cat_vars, drop_first=True)
        X_vif = df_dummy.drop(columns=[dependent_var])
        X_vif = X_vif.astype(float)
        X_vif = sm.add_constant(X_vif)
        
        vif_data = pd.DataFrame()
        vif_data["Feature"] = X_vif.columns
        vif_data["VIF"] = [variance_inflation_factor(X_vif.values, i) for i in range(X_vif.shape[1])]
        high_vif = vif_data[(vif_data["VIF"] > vif_threshold) & (vif_data["Feature"] != "const")]
        
        if not high_vif.empty:
            print("Multicollinearity detected. The following variables have VIF > {}:".format(vif_threshold))
            print(tabulate(vif_data, headers="keys", showindex=False))
            print("\n")
        else:
            print("No multicollinearity detected (all VIF values ≤ {}).\n".format(vif_threshold))
        
            # Prepare data
            X = sm.add_constant(df_test[independent_var])  # Add intercept term        
            if pd.notna(dependent_assign):
                df_test[dependent_var] = df_test[dependent_var].apply(lambda x: 1 if x == dependent_assign else 0).astype(int)
                y = df_test[dependent_var]
            else:
                y = df_test[dependent_var].cat.codes  # Convert categorical to numerical codes
            
            
            if len(independent_var) == 1:
                var = independent_var[0]
                if independent_assign and var in independent_assign:
                    independent_part = f"C({var}, Treatment(reference='{independent_assign[var]}'))"
                else:
                    independent_part = var
                
                formula = f"{dependent_var} ~ {independent_part}"
                
                model = smf.logit(formula, df_test).fit(disp=0)
                summary = model.summary().tables[1]
                summary_df = pd.DataFrame(summary)
            
            else:    
                
                summary_df = pd.DataFrame()
                
                # Univariate analysis
                for var in independent_var:
                    
                    if independent_assign and var in independent_assign:
                        independent_part = f"C({var}, Treatment(reference='{independent_assign[var]}'))"
                    else:
                        independent_part = var
                
                    formula = f"{dependent_var} ~ {independent_part}"
                    
                    model = smf.logit(formula, df_test).fit(disp=0)
                    
                    # Create a clean DataFrame
                    result_df = pd.DataFrame({
                        "predictor": model.params.index,  # Variable names
                        "coef (log OR)": np.round(model.params.values, 2),  # Log-odds (log(OR))
                        "OR": np.round(np.exp(model.params.values), 2),  # Exponentiated OR
                        "P>|z|": np.round(model.pvalues.values, 3),  # P-values formatted to 2 decimals
                        "[0.025": np.round(np.exp(model.conf_int()[0]).values, 2),  # Lower bound of CI in OR
                        "0.975]": np.round(np.exp(model.conf_int()[1]).values, 2)   # Upper bound of CI in OR
                    })
                    result_df = result_df[result_df["predictor"] != "Intercept"]
                    
                    # Append results to summary_df
                    summary_df = pd.concat([summary_df, result_df], ignore_index=True)
                

                # Multivariate analysis
                
                # Construct formula for logistic regression
                formula_parts = []    
                for var in independent_var:
                    if independent_assign and var in independent_assign:  # If variable has a reference group
                        formula_parts.append(f"C({var}, Treatment(reference='{independent_assign[var]}'))")
                    else:  # Otherwise, include normally
                        formula_parts.append(var)
                multivar_formula = f"{dependent_var} ~ " + " + ".join(formula_parts)  # Ensure no extra '+'
                
                model = smf.logit(multivar_formula, df_test).fit(disp=0)
                print(model.summary())
                
                result_df = pd.DataFrame({
                    "predictor": model.params.index,  # Variable names
                    "coef (log OR)": np.round(model.params.values, 2),  # Log-odds (log(OR))
                    "OR": np.round(np.exp(model.params.values), 2),  # Exponentiated OR
                    "P>|z|": np.round(model.pvalues.values, 3),  # P-values formatted to 2 decimals
                    "[0.025": np.round(np.exp(model.conf_int()[0]).values, 2),  # Lower bound of CI in OR
                    "0.975]": np.round(np.exp(model.conf_int()[1]).values, 2)   # Upper bound of CI in OR
                })
                result_df = result_df[result_df["predictor"] != "Intercept"]
                    
                summary_df = pd.merge(summary_df, result_df, how='left', on='predictor')
                
            print(tabulate(summary_df, showindex=False, headers="keys"))
            if export_result == True:
                ResultExport.add_to_docx(summary_df, table_name, output_dir='output/analyse')
        
        # coef = summary_df.loc[independent_var, 'Coef.']
        # conf_25 = summary_df.loc[independent_var, '[0.025']  # Lower bound of 95% CI
        # conf_975 = summary_df.loc[independent_var, '0.975]']  # Upper bound of 95% CI
        # significant_coeffs = (((conf_25 < 0) & (conf_975 > 0)) | ((conf_25 > 0) & (conf_975 < 0)))

        # # Determine if binary or multinomial logistic regression is required
        # if len(df_test[dependent_var].cat.categories) == 2:
        #     test = "Binary Logistic Regression"
        #     model = sm.Logit(y, X)
        # elif (len(df_test[dependent_var].cat.categories) > 2) & (len(df_test[dependent_var].cat.categories) <= 20):
        #     test = "Multinomial Logistic Regression"
        #     model = sm.MNLogit(y, X)
        # else:
        #     model = None
        #     print(f"Dependent variable ({dependent_var}) must have at least two categories and no more than 20 categories.")
                
        # if model is not None:
            
        #     # Fit the model
        #     fit = model.fit(disp=False)

        #     # Extract coefficients, p-values, and confidence intervals
        #     summary = fit.summary2().tables[1]
        #     summary_df = pd.DataFrame(summary)
                       
        #     if test == "Binary Logistic Regression":
        #         p_value = summary_df.loc[independent_var, 'P>|z|']  # Use 'P>|z|' for binary logistic regression
        #     elif test == "Multinomial Logistic Regression":
        #         p_value = summary_df.loc[independent_var, 'P>|t|']  # Use 'P>|t|' for multinomiallogistic regression
            

        #     if significant_coeffs:
        #         effect_size = None
        #         interpretation = 'Negligible'
        #     else:
        #         effect_size = coef
        #         interpretation = Difference.interpret_effect_size(effect_size, 'odd')

        #     if export_as == 'parameters':
        #         return p_value, effect_size, test, interpretation
        #     elif export_as == 'table':
        #         return summary_df
        #     else:
        #         raise ValueError("Invalid export_as value. Use 'parameters' or 'table'.")
        
        # else:
        #     p_value = None
        #     effect_size = None
        #     test = "Not performed any test due to number of categories"
        #     interpretation = None
        #     return p_value, effect_size, test, interpretation


# def cronbach_alpha(df):
    
#     # Number of items
#     n_items = df.shape[1]
    
#     # Variance of each item
#     item_variances = df.var(axis=0, ddof=1)
    
#     # Variance of the total score
#     total_score_variance = df.sum(axis=1).var(ddof=1)
    
#     # Cronbach's alpha formula
#     alpha = (n_items / (n_items - 1)) * (1 - (item_variances.sum() / total_score_variance))
    
#     return alpha


def logistic_regression(df_input, outcome_column, outcome_choice, predictors, reference_group=None, multivariate=False):
    
    df = df_input.copy()
    df = df[predictors + [outcome_column]]

    table_name = f"logis_{outcome_choice}_vs_others"
    print(f"Logistic regression analysis for {outcome_choice} as outcome")

    # Create a binary outcome based on the user's choice
    df['outcome'] = df[outcome_column].apply(lambda x: 1 if x == outcome_choice else 0)

    # Ensure predictors are numeric and handle categorical data
    categorical_predictors = df[predictors].select_dtypes(include=['object', 'category']).columns.tolist()
    
    # Handling reference group before creating dummies
    if reference_group:
        for predictor, ref_value in reference_group.items():
            if predictor in categorical_predictors:
                df[predictor] = pd.Categorical(df[predictor], categories=[ref_value] + [x for x in df[predictor].unique() if x != ref_value], ordered=True)     

    df = pd.get_dummies(df, columns=categorical_predictors, drop_first=True)
    df = df.applymap(lambda x: int(x) if isinstance(x, bool) else x)

    # Determine final list of predictors including dummy variables
    predictors_dummies = []
    for predictor in predictors:
        if predictor in categorical_predictors:
            
            ref_value = reference_group.get(predictor)
            # Add a reference row for the original category
            ref_row = {
                "Variables": f"{predictor}_{ref_value}",
                "u-var OR": "reference",
                "u-var 95% CI" : "",
                "p-value": "",
                "m-var OR": "reference",
                "m-var 95% CI": "",
                "m-var p-value": ""
            }
            predictors_dummies.append(ref_row)
            # Add all dummy columns corresponding to this categorical predictor
            dummies = df.columns[df.columns.str.startswith(predictor)]
            predictors_dummies.extend(dummies)
        else:
            predictors_dummies.append(predictor)
    
    results = []

    warnings.filterwarnings("ignore", category=ConvergenceWarning)

    # Univariate analysis
    for predictor in predictors_dummies:
        if isinstance(predictor, str):  # Exclude the reference row
            X = df[[predictor]]
            X = sm.add_constant(X)
            y = df['outcome']
            try:
                model = sm.Logit(y, X)
                result = model.fit(disp=0)  # Suppress output
                
                coef = np.exp(result.params[predictor])
                conf = np.exp(result.conf_int().loc[predictor])
                pvalue = result.pvalues[predictor]
                ci = f"{conf[0]:.2f}, {conf[1]:.2f}"
                row = {
                    "Variables": predictor,
                    "u-var OR": f"{coef:.2f}",
                    "u-var 95% CI" : ci,
                    "p-value": f"{pvalue:.3f}",
                    "m-var OR": "",
                    "m-var 95% CI": "",
                    "m-var p-value": ""
                }
                results.append(row)
            except Exception as e:
                print(f"Error fitting univariate model for {predictor}: {e}")
                continue
        else:
            # Append reference row as is
            results.append(predictor)

    # Multivariate analysis
    if multivariate:
        X = df[[pred for pred in predictors_dummies if isinstance(pred, str)]]
        X = sm.add_constant(X)
        y = df['outcome']
        try:
            model = sm.Logit(y, X)
            result = model.fit(disp=0)  # Suppress output

            for predictor in predictors_dummies:
                if isinstance(predictor, str):  # Exclude the reference row
                    coef = np.exp(result.params[predictor])
                    conf = np.exp(result.conf_int().loc[predictor])
                    pvalue = result.pvalues[predictor]
                    ci = f"{conf[0]:.2f}, {conf[1]:.2f}"
                    
                    # Find the row in the results and update it
                    for row in results:
                        if row["Variables"] == predictor:
                            row["m-var OR"] = f"{coef:.2f}"
                            row['m-var 95% CI'] = ci
                            row["m-var p-value"] = f"{pvalue:.3f}"
                            break
        except Exception as e:
            print(f"Error fitting multivariate model: {e}")

    # Display the results
    print(tabulate(results, headers="keys"))
    add_to_docx(results, table_name, output_dir='output/analyse')

# import pandas as pd
# import statsmodels.api as sm
# import matplotlib.pyplot as plt

# import pandas as pd
# import statsmodels.api as sm
# import matplotlib.pyplot as plt

# def interrupted_time_series(df_input, time_column, outcome_column=None, intervention_point=None, 
#                             control_columns=None, time_unit='date', show_summary=True, plot=True, ax=None, point_size=None, 
#                             title='Title', axis_label_size=14, title_size=16, tick_label_size=12, line_width=2, y_lim=None, first_col=False,
#                             counterfactual_line=True):
#     """
#     Perform Interrupted Time Series analysis.
    
#     Parameters:
#     - df: pd.DataFrame, the data containing the time series.
#     - time_column: str, the column name representing time.
#     - outcome_column: str, optional, the column name representing the outcome variable. If None, use counts.
#     - intervention_point: int/str/datetime, optional, the time point at which the intervention occurred.
#     - control_columns: list of str, optional, names of control variables.
#     - time_unit: str, optional, the unit of time ('minute', 'hour', 'day', 'month', etc.). Default is 'day'.
#     - plot: bool, optional, whether to plot the time series and regression lines.
    
#     Returns:
#     - results: Regression results summary from statsmodels.
#     """

#     df = df_input.copy()

#     # Convert the time column to a datetime format
#     df[time_column] = pd.to_datetime(df[time_column], errors='coerce')
    
#     # Check for any non-datetime values and handle them (e.g., drop or fill with a default value)
#     if df[time_column].isnull().any():
#         print(f"Warning: {df[time_column].isnull().sum()} non-datetime entries found in {time_column}. These will be dropped.")
#         df = df.dropna(subset=[time_column])
    
#     # Manipulate the time column based on the specified time unit
#     if time_unit == 'date':
#         df[time_column] = df[time_column].dt.floor('D')  # Only keep the date part
#     elif time_unit == 'month':
#         df[time_column] = df[time_column].dt.to_period('M').dt.to_timestamp()  # Keep only month and year
#     elif time_unit == 'year':
#         df[time_column] = df[time_column].dt.to_period('Y').dt.to_timestamp()  # Keep only the year
#     # For 'minute', 'hour', 'week', and other time units, no further manipulation is needed
    
#     if outcome_column is None:
#         df = df.groupby(time_column).size().reset_index(name='outcome')
#         outcome_column = 'outcome'
#     else:
#         df = df.groupby(time_column).agg({outcome_column: 'mean'}).reset_index()
    
#     # Create time variable based on the specified time unit
#     time_conversion = {
#         'minute': 'T',
#         'hour': 'H',
#         'date': 'D',
#         'week': 'W',
#         'month': 'M',
#         'year': 'Y'
#     }
#     df['time'] = df[time_column].dt.to_period(time_conversion.get(time_unit, 'D')).astype(str)
#     df['time'] = pd.to_datetime(df['time']).rank(method='first').astype(int)

#     # Set the intervention point
#     if intervention_point is not None:
#         if isinstance(intervention_point, str) or isinstance(intervention_point, pd.Timestamp):
#             intervention_point = df[df[time_column] >= pd.to_datetime(intervention_point)].iloc[0]['time']

#     # Create pre- and post-intervention indicators
#     df['intervention'] = (df['time'] >= intervention_point).astype(int) if intervention_point else 0
    
#     # Create time after intervention variable
#     df['time_after_intervention'] = df['time'] - df['time'][df['intervention'] == 1].min()
#     df['time_after_intervention'] = df['time_after_intervention'].apply(lambda x: x if x >= 0 else 0)
    
#     # Create design matrix
#     X = sm.add_constant(df[['time', 'intervention', 'time_after_intervention']])
    
#     # Include control variables if provided
#     if control_columns:
#         X = sm.add_constant(df[['time', 'intervention', 'time_after_intervention'] + control_columns])
    
#     # Fit the model
#     model = sm.OLS(df[outcome_column], X)
#     results = model.fit()
    
#     if show_summary:
#         print(results.summary())
    
#     if plot:
        
#         if ax is None:
#             fig, ax = plt.subplots(figsize=(10, 6))
        
#         # Plot the time series with intervention
#         ax.scatter(df['time'], df[outcome_column], label='Outcome', color='grey', alpha=0.5, s=point_size)
        
#         if intervention_point:
#             ax.axvline(x=df['time'][df['time'] == intervention_point].iloc[0], color='red', linestyle='--', label='Intervention Point', linewidth=line_width)
        
#         if counterfactual_line and intervention_point:
#             # Create the counterfactual line by extending the pre-intervention trend
#             pre_intervention_model = sm.OLS(df[outcome_column][df['time'] < intervention_point], 
#                                             sm.add_constant(df[['time']][df['time'] < intervention_point])).fit()
#             df['counterfactual'] = pre_intervention_model.predict(sm.add_constant(df[['time']]))
#             ax.plot(df['time'], df['counterfactual'], label='Counterfactual', color='green', linestyle='--', linewidth=line_width)
        
#         df['predicted'] = results.predict(X)
#         ax.plot(df['time'], df['predicted'], label='Fitted values', color='blue', linewidth=line_width)
#         if y_lim is not None:  # Only set y_lim if it is provided
#             ax.set_ylim(0, y_lim)
#         ax.set_xlabel(time_unit, fontsize=axis_label_size)
#         if first_col:
#             ax.set_ylabel(outcome_column, fontsize=axis_label_size)
#         ax.set_title(title, fontsize=title_size)
#         ax.tick_params(axis='both', labelsize=tick_label_size)
#         ax.grid(True)
    
#     return results


# def interrupted_time_series_with_counterfactual(df_input, time_column, outcome_column=None, intervention_point=None, 
#                                                                   split_point=None, control_columns=None, time_unit='date', 
#                                                                   show_summary=True, plot=True, ax=None, point_size=10, 
#                                                                   title='Title', axis_label_size=14, title_size=16, 
#                                                                   tick_label_size=12, line_width=2, y_lim=(0,1), first_col=False,
#                                                                   counterfactual_postanalysis=True):
#     """
#     Perform Interrupted Time Series analysis with a split for counterfactual comparison and additional post-intervention regression.
    
#     Parameters:
#     - df_input: pd.DataFrame, the data containing the time series.
#     - time_column: str, the column name representing time.
#     - outcome_column: str, optional, the column name representing the outcome variable. If None, use counts.
#     - intervention_point: int/str/datetime, optional, the time point at which the intervention occurred.
#     - split_point: int/str/datetime, optional, the time point at which to split the data for counterfactual comparison.
#     - control_columns: list of str, optional, names of control variables.
#     - time_unit: str, optional, the unit of time ('minute', 'hour', 'day', 'month', etc.). Default is 'day'.
#     - plot: bool, optional, whether to plot the time series and regression lines.
#     - counterfactual_line: bool, optional, whether to plot the counterfactual (pre-intervention trend) in the post-intervention period.

#     Returns:
#     - results_pre_split: Regression results summary from pre-split analysis.
#     - results_post_split: Regression results summary from post-split analysis.
#     """

#     df = df_input.copy()

#     # Convert the time column to a datetime format
#     df[time_column] = pd.to_datetime(df[time_column], errors='coerce')
    
#     # Check for any non-datetime values and handle them (e.g., drop or fill with a default value)
#     if df[time_column].isnull().any():
#         print(f"Warning: {df[time_column].isnull().sum()} non-datetime entries found in {time_column}. These will be dropped.")
#         df = df.dropna(subset=[time_column])
    
#     # Manipulate the time column based on the specified time unit
#     if time_unit == 'date':
#         df[time_column] = df[time_column].dt.floor('D')  # Only keep the date part
#     elif time_unit == 'month':
#         df[time_column] = df[time_column].dt.to_period('M').dt.to_timestamp()  # Keep only month and year
#     elif time_unit == 'year':
#         df[time_column] = df[time_column].dt.to_period('Y').dt.to_timestamp()  # Keep only the year
#     # For 'minute', 'hour', 'week', and other time units, no further manipulation is needed
    
#     if outcome_column is None:
#         df = df.groupby(time_column).size().reset_index(name='outcome')
#         outcome_column = 'outcome'
#     else:
#         df = df.groupby(time_column).agg({outcome_column: 'mean'}).reset_index()
    
#     # Create time variable based on the specified time unit
#     time_conversion = {
#         'minute': 'T',
#         'hour': 'H',
#         'date': 'D',
#         'week': 'W',
#         'month': 'M',
#         'year': 'Y'
#     }
#     # Get the maximum and minimum time
#     max_time = df[time_column].max()
#     min_time = df[time_column].min()

#     # Create the time range DataFrame
#     its_time_range = pd.date_range(start=min_time, end=max_time, freq=time_conversion.get(time_unit, 'D'))
#     its_df = pd.DataFrame({time_column: its_time_range})
#     its_df['time'] = its_df[time_column].dt.to_period(time_conversion.get(time_unit, 'D')).astype(str)
#     its_df['time'] = pd.to_datetime(its_df['time']).rank(method='first').astype(int)
    
#     its_df = pd.merge(its_df, df[[time_column, outcome_column]], how='left', on=time_column)

#     # Set the intervention point
#     if intervention_point is not None:
#         if isinstance(intervention_point, str) or isinstance(intervention_point, pd.Timestamp):
#             intervention_point = its_df[its_df[time_column] >= pd.to_datetime(intervention_point)].iloc[0]['time']

#     # Set the split points
#     if split_point is not None:
#         if isinstance(split_point, (tuple, list)) and len(split_point) == 2:
#             start_gap = pd.to_datetime(split_point[0])
#             end_gap = pd.to_datetime(split_point[1])
#         else:
#             raise ValueError("split_point must be a tuple or list with two elements: (start_gap, end_gap)")
    
#     # Split the data into pre-split and post-split groups
#     df_pre_split = its_df[its_df[time_column] <= start_gap].copy()
#     counterfactual_split = its_df[its_df[time_column] > start_gap].copy()
#     df_post_split = its_df[its_df[time_column] >= end_gap].copy()
    
#     # Perform ITS on the pre-split data
#     df_pre_split['intervention'] = (df_pre_split['time'] >= intervention_point).astype(int) if intervention_point else 0
#     df_pre_split['time_after_intervention'] = df_pre_split['time'] - df_pre_split['time'][df_pre_split['intervention'] == 1].min()
#     df_pre_split['time_after_intervention'] = df_pre_split['time_after_intervention'].apply(lambda x: x if x >= 0 else 0)

#     # Create design matrix for pre-split data
#     X_pre_split = sm.add_constant(df_pre_split[['time', 'intervention', 'time_after_intervention']])
    
#     # Include control variables if provided
#     if control_columns:
#         X_pre_split = sm.add_constant(df_pre_split[['time', 'intervention', 'time_after_intervention'] + control_columns])
    
#     # Fit the model on pre-split data
#     model_pre_split = sm.OLS(df_pre_split[outcome_column], X_pre_split)
#     results_pre_split = model_pre_split.fit()
    
#     if show_summary:
#         print("ITS Analysis for Pre-Split Data:")
#         print(results_pre_split.summary())
    
#     # Calculate the predicted values for the pre-split data
#     df_pre_split['predicted'] = results_pre_split.predict(X_pre_split)
    
#     # Predict the counterfactual trend for the post-split data using pre-split model
#     if counterfactual_postanalysis:
#         # Prepare the counterfactual data
#         counterfactual_split['intervention'] = 1  # No intervention is considered in the counterfactual scenario
#         counterfactual_split['time_after_intervention'] = counterfactual_split['time'] - intervention_point
#         counterfactual_split['time_after_intervention'] = counterfactual_split['time_after_intervention'].apply(lambda x: x if x >= 0 else 0)
        
#         b0 = results_pre_split.params['const']
#         b1 = results_pre_split.params['time']
#         b2 = results_pre_split.params['intervention']
#         b3 = results_pre_split.params['time_after_intervention']

#         counterfactual_split['counterfactual'] = b0 + (b1 * counterfactual_split['time']) + (b2) + (b3 * counterfactual_split['time_after_intervention'])
    
#     if not df_post_split.empty:
        
#         # Rank the time column to ensure it's in an integer format
#         df_post_split['time'] = pd.to_datetime(df_post_split['time']).rank(method='first').astype(int)
#         df_post_split[outcome_column].fillna(0, inplace=True)
#         X_post_split = sm.add_constant(df_post_split[['time']])
        
#         # Fit the linear regression model
#         model_post_split = sm.OLS(df_post_split[outcome_column], X_post_split).fit()
                
#         # Generate the predicted values from the linear regression model
#         df_post_split['predicted_post'] = model_post_split.fittedvalues
#         df_post_split['time'] = df_post_split[time_column].apply(lambda x: (x.date() - min_time.date()).days)
    
#     if plot:
        
#         if ax is None:
#             fig, ax = plt.subplots(figsize=(10, 6))
        
#         # Plot the actual outcome as a scatter plot with grey color and 50% opacity
#         ax.scatter(its_df['time'], its_df[outcome_column], label='Outcome', color='grey', alpha=0.5, s=point_size)
        
#         if intervention_point:
#             ax.axvline(x=its_df['time'][its_df['time'] == intervention_point].iloc[0], color='red', linestyle='--', label='Intervention Point', linewidth=line_width)
        
#         # Plot the fitted values from the pre-split data
#         ax.plot(df_pre_split['time'], df_pre_split['predicted'], label='Fitted values (Pre-Split)', color='blue', linewidth=line_width)
        
#         if counterfactual_postanalysis and not df_post_split.empty:
#             # Plot the counterfactual line for post-split data
#             ax.plot(counterfactual_split['time'], counterfactual_split['counterfactual'], label='Counterfactual (Post-Analysis)', color='green', linestyle='--', linewidth=line_width)
        
#         if not df_post_split.empty:
#             # Plot the post-split regression line
#             ax.plot(df_post_split['time'], df_post_split['predicted_post'], label='Fitted values (Post-Split)', color='purple', linewidth=line_width)
        
#         ax.set_ylim(y_lim)
        
#         # Set axis labels and title with specified font sizes
#         ax.set_xlabel(time_unit, fontsize=axis_label_size)
#         if first_col:
#             ax.set_ylabel(outcome_column, fontsize=axis_label_size)
#         ax.set_title(title, fontsize=title_size)
        
#         # Set the size of the tick labels
#         ax.tick_params(axis='x', labelsize=tick_label_size)
#         ax.tick_params(axis='y', labelsize=tick_label_size)
        
#         ax.grid(True)


#     return results_pre_split







# from sklearn.ensemble import RandomForestClassifier
# from sklearn.metrics import classification_report, confusion_matrix, roc_auc_score
# from sklearn.model_selection import train_test_split
# from sklearn.preprocessing import LabelEncoder
# import pandas as pd
# import numpy as np
# import warnings

# def encode_categorical_columns(df, columns):
#     """Encode categorical columns with label encoding."""
#     label_encoders = {}
#     for col in columns:
#         le = LabelEncoder()
#         df[col] = le.fit_transform(df[col].astype(str))  # Ensure all data is string before encoding
#         label_encoders[col] = le
#     return df, label_encoders

# def random_forest(df_input, outcome_column, outcome_choice, predictors, reference_group=None, test_size=0.2, random_state=42):
    
#     df = df_input.copy()
#     df = df[predictors + [outcome_column]]

#     table_name = f"random_forest_{outcome_choice}_vs_others"
#     print(f"Random Forest analysis for {outcome_choice} as outcome")

#     # Encode the outcome column
#     le_outcome = LabelEncoder()
#     df['outcome'] = le_outcome.fit_transform(df[outcome_column].astype(str))

#     # Print classes in the outcome for reference
#     print(f"Outcome encoding: {dict(zip(le_outcome.classes_, le_outcome.transform(le_outcome.classes_)))}")

#     # Ensure predictors are numeric and handle categorical data
#     categorical_predictors = df[predictors].select_dtypes(include=['object', 'category']).columns.tolist()

#     # Handling reference group before encoding
#     if reference_group:
#         for predictor, ref_value in reference_group.items():
#             if predictor in categorical_predictors:
#                 df[predictor] = pd.Categorical(df[predictor], categories=[ref_value] + [x for x in df[predictor].unique() if x != ref_value], ordered=True)     

#     # Encode categorical predictors
#     df, _ = encode_categorical_columns(df, categorical_predictors)

#     X = df.drop(columns=[outcome_column, 'outcome'])
#     y = df['outcome']

#     # Split the data into training and test sets
#     X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=test_size, random_state=random_state)

#     # Initialize and train the Random Forest model
#     model = RandomForestClassifier(random_state=random_state)
#     model.fit(X_train, y_train)

#     # Predictions
#     y_pred = model.predict(X_test)
#     y_pred_proba = model.predict_proba(X_test)[:, 1]

#     # Evaluation metrics
#     print("\nClassification Report:")
#     print(classification_report(y_test, y_pred))
    
#     print("\nConfusion Matrix:")
#     print(confusion_matrix(y_test, y_pred))

#     roc_auc = roc_auc_score(y_test, y_pred_proba)
#     print(f"\nROC AUC Score: {roc_auc:.3f}")

#     # Feature importance
#     importances = model.feature_importances_
#     feature_importance = pd.DataFrame({'Feature': X.columns, 'Importance': importances})
#     feature_importance = feature_importance.sort_values(by='Importance', ascending=False)

#     print("\nFeature Importance:")
#     print(feature_importance)
    
#     # Optionally save the results
#     # add_to_docx(feature_importance, table_name, output_dir='output/analyse')

# # Example usage
# # random_forest_analysis(df_logis, outcome_column, outcome_choice, predictors, reference_group)

    
#     # Optionally save the results
#     # add_to_docx(feature_importance, table_name, output_dir='output/analyse')

# # Example usage
# # random_forest_analysis(df_logis, outcome_column, outcome_choice, predictors, reference_group)
